"""Generate a per-grant Word report for GrantFlow."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# Color constants
NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(11, 76, 140)
MUTED = RGBColor(88, 97, 116)
GREEN = RGBColor(29, 111, 184)
AMBER = RGBColor(232, 169, 22)
RED = RGBColor(192, 39, 45)

# Status to color mapping
STATUS_COLORS = {
    "Paid": GREEN,
    "Submitted": GREEN,
    "Received": GREEN,
    "Pending": AMBER,
    "Partial": AMBER,
    "Overdue": RED,
    "Discrepancy": RED,
    "Not Received": RED,
}


def _status_color(value: str) -> RGBColor:
    """Return RGBColor for a status value."""
    return STATUS_COLORS.get(value, MUTED)


def _shade(cell, fill: str) -> None:
    """Apply background color to a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _font(run, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    """Apply font settings to a run of text."""
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _set_widths(table, widths: list[float]) -> None:
    """Set fixed column widths on a table."""
    table.autofit = False
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)


def _add_section_heading(document: Document, text: str) -> None:
    """Add a styled section heading paragraph."""
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(8)
    _font(heading.add_run(text), 13, bold=True, color=BLUE)


def _add_detail_table(document: Document, rows: list[tuple]) -> None:
    """Add a two-column label/value table.

    rows: list of (label, value, optional_color) tuples
    """
    table = document.add_table(rows=len(rows), cols=2)
    _set_widths(table, [2.2, 4.3])

    for index, row_data in enumerate(rows):
        label = row_data[0]
        value = str(row_data[1])
        color = row_data[2] if len(row_data) > 2 else MUTED

        # Label cell — shaded background
        label_cell = table.cell(index, 0)
        _shade(label_cell, "E8EEF5")
        _font(label_cell.paragraphs[0].add_run(label), 10, bold=True, color=NAVY)

        # Value cell
        value_cell = table.cell(index, 1)
        _font(value_cell.paragraphs[0].add_run(value), 10, color=color)


def create_word_report(grant: dict, output_path: Path) -> Path:
    """Generate a Word report for a single grant.

    grant: dictionary with snake_case keys from the sheet
    output_path: where to save the .docx file
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()

    # Page margins
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.75)
    section.header_distance = section.footer_distance = Inches(0.49)

    # Default font
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)

    # ── Page Header ────────────────────────────────────────────────────────
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font(
        header.add_run("IHHN GrantFlow | Grant Utilization Report"),
        8.5, color=MUTED
    )

    # ── Title Block ────────────────────────────────────────────────────────
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    _font(title.add_run("IHHN GrantFlow"), 24, bold=True, color=NAVY)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(2)
    _font(subtitle.add_run("Grant Utilization Report"), 14, bold=True, color=BLUE)

    generated = document.add_paragraph()
    generated.paragraph_format.space_after = Pt(16)
    _font(
        generated.add_run(f"Generated: {date.today().isoformat()}"),
        9, color=MUTED
    )

    # ── Grant Details Section ──────────────────────────────────────────────
    _add_section_heading(document, "Grant Details")
    _add_detail_table(document, [
        ("Grant Number", grant.get("grant_number", "—")),
        ("Supplier",     grant.get("supplier", "—")),
        ("Item",         grant.get("item", "—")),
        ("Department",   grant.get("department", "—")),
        ("Currency",     grant.get("currency", "—")),
    ])

    # ── Payment Section ────────────────────────────────────────────────────
    _add_section_heading(document, "Payment Summary")
    payment_status = grant.get("payment_status", "—")

    try:
        payment_made = float(grant.get("payment_made", 0))
        payment_usd = float(grant.get("payment_made_usd", 0))
    except (ValueError, TypeError):
        payment_made = 0.0
        payment_usd = 0.0

    _add_detail_table(document, [
        (
            "Payment Made",
            f"{grant.get('currency', '')} {payment_made:,.2f}"
        ),
        (
            "Payment Made (USD)",
            f"$ {payment_usd:,.2f}"
        ),
        (
            "Payment Status",
            payment_status,
            _status_color(payment_status)
        ),
    ])

    # ── Status Section ─────────────────────────────────────────────────────
    _add_section_heading(document, "Status Overview")
    report_status = grant.get("report_status", "—")
    shipping = grant.get("shipping", "—")
    receiving = grant.get("receiving", "—")

    _add_detail_table(document, [
        ("Report Status", report_status, _status_color(report_status)),
        ("Shipping",      shipping,      _status_color(shipping)),
        ("Receiving",     receiving,     _status_color(receiving)),
    ])

    # ── Footer Note ────────────────────────────────────────────────────────
    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(24)
    _font(
        note.add_run("Confidential — For internal use only."),
        9, color=MUTED
    )

    # Page footer
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(
        footer.add_run("IHHN GrantFlow | Internal Use Only"),
        8, color=MUTED
    )

    document.save(output_path)
    return output_path

